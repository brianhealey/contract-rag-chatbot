"""
Multi-format document parser supporting PDF, HTML, TXT, DOCX, and more.
"""

from typing import List, Dict, Any
from pathlib import Path
import hashlib
from dataclasses import dataclass
from datetime import datetime

# Document parsing libraries
import pdfplumber
from bs4 import BeautifulSoup
import docx

from config.settings import settings


@dataclass
class Document:
    """Represents a parsed document with metadata."""

    id: str
    text: str
    metadata: Dict[str, Any]
    source_file: str
    file_type: str
    parsed_at: str


class DocumentParser:
    """
    Parser for multiple document formats using unstructured library.
    """

    def __init__(self, documents_dir: str = None):
        """
        Initialize the document parser.

        Args:
            documents_dir: Directory containing documents
        """
        self.documents_dir = Path(documents_dir or settings.document.documents_dir)
        self.supported_formats = settings.document.supported_formats

    def _generate_document_id(self, file_path: str, content: str = None) -> str:
        """
        Generate a unique ID for a document.

        Args:
            file_path: Path to the document
            content: Optional content to include in hash

        Returns:
            MD5 hash as document ID
        """
        # Use file path and optionally content for ID generation
        id_string = f"{file_path}"
        if content:
            id_string += f":{content[:100]}"  # Include first 100 chars
        return hashlib.md5(id_string.encode()).hexdigest()

    def _get_file_type(self, file_path: Path) -> str:
        """Get the file type from extension."""
        return file_path.suffix.lstrip(".").lower()

    def _parse_pdf_with_pdfplumber(self, file_path: Path) -> str:
        """Parse PDF using pdfplumber."""
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n\n".join(text_parts)

    def _parse_html(self, file_path: Path) -> str:
        """Parse HTML using BeautifulSoup."""
        with open(file_path, "r", encoding="utf-8") as f:
            soup = BeautifulSoup(f.read(), "html.parser")
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            return soup.get_text(separator="\n\n")

    def _parse_docx(self, file_path: Path) -> str:
        """Parse DOCX using python-docx."""
        doc = docx.Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        return "\n\n".join(text_parts)

    def parse_file(self, file_path: str, extract_metadata: bool = True) -> Document:
        """
        Parse a single document file.

        Args:
            file_path: Path to the document file
            extract_metadata: Whether to extract document metadata

        Returns:
            Parsed Document object
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        file_type = self._get_file_type(file_path)
        if file_type not in self.supported_formats:
            raise ValueError(
                f"Unsupported file format: {file_type}. "
                f"Supported formats: {', '.join(self.supported_formats)}"
            )

        print(f"Parsing {file_type.upper()} file: {file_path.name}")

        # Parse based on file type
        try:
            if file_type == "pdf":
                # Use pdfplumber for reliable PDF parsing
                text = self._parse_pdf_with_pdfplumber(file_path)
            elif file_type == "html":
                text = self._parse_html(file_path)
            elif file_type == "txt" or file_type == "md":
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
            elif file_type == "docx":
                text = self._parse_docx(file_path)
            else:
                # Try reading as text
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()

        except Exception as e:
            print(f"Error parsing {file_path}: {e}")
            raise

        # Build metadata
        metadata = {
            "source_file": file_path.name,
            "file_path": str(file_path),
            "file_type": file_type,
            "file_size": file_path.stat().st_size,
            "parsed_at": datetime.now().isoformat(),
            "text_length": len(text),
        }

        # Generate document ID
        doc_id = self._generate_document_id(str(file_path), text)

        return Document(
            id=doc_id,
            text=text,
            metadata=metadata,
            source_file=file_path.name,
            file_type=file_type,
            parsed_at=datetime.now().isoformat(),
        )

    def parse_directory(
        self, directory: str = None, recursive: bool = False
    ) -> List[Document]:
        """
        Parse all supported documents in a directory.

        Args:
            directory: Directory to parse (defaults to documents_dir)
            recursive: Whether to search recursively

        Returns:
            List of parsed Document objects
        """
        directory = Path(directory or self.documents_dir)

        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {directory}")

        documents = []

        # Get all files with supported extensions
        patterns = [f"*.{ext}" for ext in self.supported_formats]

        for pattern in patterns:
            if recursive:
                files = directory.rglob(pattern)
            else:
                files = directory.glob(pattern)

            for file_path in files:
                try:
                    doc = self.parse_file(str(file_path))
                    documents.append(doc)
                except Exception as e:
                    print(f"Failed to parse {file_path}: {e}")
                    continue

        print(f"\nParsed {len(documents)} documents from {directory}")
        return documents

    def get_file_list(
        self, directory: str = None, recursive: bool = False
    ) -> List[Path]:
        """
        Get list of supported document files in a directory.

        Args:
            directory: Directory to search
            recursive: Whether to search recursively

        Returns:
            List of file paths
        """
        directory = Path(directory or self.documents_dir)
        files = []

        patterns = [f"*.{ext}" for ext in self.supported_formats]

        for pattern in patterns:
            if recursive:
                files.extend(directory.rglob(pattern))
            else:
                files.extend(directory.glob(pattern))

        return sorted(files)


def main():
    """Test the document parser."""
    parser = DocumentParser()

    print("=== Testing Document Parser ===\n")

    # Get list of files
    files = parser.get_file_list()
    print(f"Found {len(files)} document(s) in {parser.documents_dir}:\n")

    if not files:
        print("No documents found. Please add documents to the 'documents' directory.")
        print(f"Supported formats: {', '.join(parser.supported_formats)}")
        return

    for file in files:
        print(f"  - {file.name} ({file.suffix})")

    # Parse all documents
    print("\n=== Parsing Documents ===\n")
    documents = parser.parse_directory()

    # Display results
    for i, doc in enumerate(documents, 1):
        print(f"\n--- Document {i} ---")
        print(f"ID: {doc.id}")
        print(f"File: {doc.source_file}")
        print(f"Type: {doc.file_type}")
        print(f"Size: {doc.metadata.get('file_size', 0)} bytes")
        print(f"Text length: {len(doc.text)} characters")
        print(f"First 200 chars:\n{doc.text[:200]}...")


if __name__ == "__main__":
    main()
